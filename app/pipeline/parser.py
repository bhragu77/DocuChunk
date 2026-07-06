"""
Pipeline Stage 1 — Document Parsing

Extracts raw text from PDF and DOCX files.

Learning note: the output is deliberately "dirty". We keep all the raw text
including page headers, footers, hyphenated words, and extra whitespace.
That is the cleaner's job (Stage 2). Separating concerns lets you tune each
stage independently.

PDF  → PyMuPDF (fitz): page-by-page extraction, preserves page numbers.
DOCX → python-docx: paragraph + table extraction, groups into logical pages.
"""
import io
import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.oxml.ns import qn

from app.pipeline.types import ParsedPage

logger = logging.getLogger(__name__)

# How many non-empty DOCX paragraphs to group into one logical "page".
# DOCX has no native page concept, so we chunk paragraphs into groups.
_DOCX_PARAGRAPHS_PER_PAGE = 40


def parse(file_path: str, source: str, mime_type: str) -> list[ParsedPage]:
    """
    Dispatch to the correct parser based on mime_type.
    Returns a list of ParsedPage — one entry per page (PDF) or paragraph group (DOCX).
    """
    if mime_type == "application/pdf":
        return _parse_pdf(file_path, source)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(file_path, source)
    else:
        raise ValueError(f"Unsupported mime type: {mime_type}")


# ── PDF ───────────────────────────────────────────────────────────────────────

def _parse_pdf(file_path: str, source: str) -> list[ParsedPage]:
    """
    Extract text from each page of a PDF.

    Learning note: PyMuPDF's get_text("text") returns text in reading order.
    Some PDFs are scanned images — get_text returns empty string for those.
    For such pages we fall back to OCR (Tesseract via pytesseract); see _ocr_page.
    """
    pages: list[ParsedPage] = []

    try:
        doc = fitz.open(file_path)
    except Exception as e:
        raise ValueError(f"Failed to open PDF '{source}': {e}") from e

    total_pages = doc.page_count
    logger.info("Parsing PDF '%s': %d pages", source, total_pages)

    for page_num in range(total_pages):
        page = doc[page_num]
        text = page.get_text("text")  # reading-order plain text

        if not text.strip():
            # Scanned/image-only page: no text layer. Fall back to OCR before
            # giving up, so scanned PDFs still make it through the pipeline.
            logger.info("Page %d has no text layer — attempting OCR", page_num + 1)
            text = _ocr_page(page, page_num + 1)

        if not text.strip():
            logger.debug("Page %d is blank or OCR yielded nothing — skipping", page_num + 1)
            continue

        pages.append(ParsedPage(
            text=text,
            page_number=page_num + 1,  # 1-indexed
            source=source,
        ))

    doc.close()
    logger.info("PDF parsed: %d non-blank pages extracted", len(pages))
    return pages


# OCR resolution. 300 DPI is the accuracy/speed sweet spot for Tesseract on
# document scans — lower loses small text, higher mostly just costs time.
_OCR_DPI = 300


def _ocr_page(page, page_num: int) -> str:
    """
    OCR fallback for an image-only PDF page.

    Renders the page to a raster image and runs Tesseract over it. Returns the
    extracted text (possibly empty). Best-effort: if the OCR stack isn't available
    (pytesseract/Pillow not installed, or the tesseract binary missing) or OCR
    fails for any reason, we log and return "" so the page is simply skipped —
    the pipeline degrades to its old behaviour rather than crashing.
    """
    try:
        import pytesseract
        from PIL import Image
    except Exception as e:
        logger.warning("OCR unavailable (pytesseract/Pillow not importable): %s", e)
        return ""

    try:
        pix = page.get_pixmap(dpi=_OCR_DPI)
        # Round-trip through PNG bytes so PIL gets the colourspace/alpha right
        # regardless of the source page's pixel format.
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(img)
        logger.info("OCR page %d: extracted %d chars", page_num, len(text.strip()))
        return text
    except Exception as e:
        logger.warning("OCR failed for page %d: %s", page_num, e)
        return ""


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _parse_docx(file_path: str, source: str) -> list[ParsedPage]:
    """
    Extract text from a DOCX file: body paragraphs + tables.

    Learning note: DOCX has no native page boundary data at the paragraph level.
    We simulate pages by grouping every N paragraphs together. This is a common
    approximation — real page detection would require rendering the document.
    """
    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        raise ValueError(f"Failed to open DOCX '{source}': {e}") from e

    pages: list[ParsedPage] = []
    page_number = 1

    # ── Body paragraphs ──
    para_buffer: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        para_buffer.append(text)

        if len(para_buffer) >= _DOCX_PARAGRAPHS_PER_PAGE:
            pages.append(ParsedPage(
                text="\n\n".join(para_buffer),
                page_number=page_number,
                source=source,
            ))
            page_number += 1
            para_buffer = []

    # Flush any remaining paragraphs
    if para_buffer:
        pages.append(ParsedPage(
            text="\n\n".join(para_buffer),
            page_number=page_number,
            source=source,
        ))
        page_number += 1

    # ── Tables ──
    # Tables are separate from paragraphs in DOCX's XML structure.
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))

        if rows:
            pages.append(ParsedPage(
                text="\n".join(rows),
                page_number=page_number,
                source=source,
            ))
            page_number += 1

    logger.info("DOCX parsed: %d logical pages extracted from '%s'", len(pages), source)
    return pages
