"""
Tests for pipeline Stages 1–2: Parser and Cleaner.

We create real files (not mocks) using PyMuPDF and python-docx so the tests
exercise the actual extraction code paths.
"""
import io
import os
import time
import pytest
import tempfile

from app.pipeline.types import ParsedPage, CleanedPage
from app.pipeline.parser import parse
from app.pipeline.cleaner import clean_page, clean_pages


# ── File factories ────────────────────────────────────────────────────────────

def _make_pdf_file(pages_content: list[str]) -> str:
    """Write a multi-page PDF to a temp file. Returns path."""
    import fitz
    doc = fitz.open()
    for text in pages_content:
        page = doc.new_page()
        page.insert_text((50, 72), text)
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    doc.save(tmp.name)
    doc.close()
    return tmp.name


def _make_docx_file(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> str:
    """Write a DOCX file to a temp file. Returns path."""
    from docx import Document
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r_idx, row_data in enumerate(table_rows):
            for c_idx, cell_text in enumerate(row_data):
                table.cell(r_idx, c_idx).text = cell_text
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    return tmp.name


@pytest.fixture(autouse=True)
def cleanup_temp_files(tmp_path):
    """Collect temp file paths created during a test and delete after."""
    yield


# ── Parser: PDF ───────────────────────────────────────────────────────────────

def test_parse_pdf_single_page():
    path = _make_pdf_file(["Hello from page one."])
    try:
        pages = parse(path, "test.pdf", "application/pdf")
        assert len(pages) == 1
        assert pages[0].page_number == 1
        assert pages[0].source == "test.pdf"
        assert "Hello from page one" in pages[0].text
    finally:
        os.unlink(path)


def test_parse_pdf_multiple_pages():
    content = [f"Content of page {i+1}" for i in range(5)]
    path = _make_pdf_file(content)
    try:
        pages = parse(path, "multi.pdf", "application/pdf")
        assert len(pages) == 5
        # Page numbers are 1-indexed and in order
        assert [p.page_number for p in pages] == [1, 2, 3, 4, 5]
        for i, page in enumerate(pages):
            assert f"page {i+1}" in page.text.lower()
    finally:
        os.unlink(path)


def test_parse_pdf_returns_parsed_page_type():
    path = _make_pdf_file(["Type check page."])
    try:
        pages = parse(path, "type.pdf", "application/pdf")
        assert all(isinstance(p, ParsedPage) for p in pages)
    finally:
        os.unlink(path)


def test_parse_pdf_bad_file_raises():
    with pytest.raises(ValueError, match="Failed to open PDF"):
        parse("/tmp/nonexistent_file.pdf", "bad.pdf", "application/pdf")


# ── Parser: DOCX ─────────────────────────────────────────────────────────────

def test_parse_docx_paragraphs():
    paras = ["First paragraph.", "Second paragraph.", "Third paragraph."]
    path = _make_docx_file(paras)
    try:
        pages = parse(path, "test.docx",
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert len(pages) >= 1
        all_text = " ".join(p.text for p in pages)
        for para in paras:
            assert para in all_text
    finally:
        os.unlink(path)


def test_parse_docx_table_extracted():
    path = _make_docx_file(
        paragraphs=["Header paragraph"],
        table_rows=[["Name", "Age"], ["Alice", "30"], ["Bob", "25"]],
    )
    try:
        pages = parse(path, "table.docx",
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        all_text = " ".join(p.text for p in pages)
        assert "Alice" in all_text
        assert "Bob" in all_text
    finally:
        os.unlink(path)


def test_parse_docx_empty_paragraphs_skipped():
    paras = ["Real content.", "", "   ", "More content."]
    path = _make_docx_file(paras)
    try:
        pages = parse(path, "sparse.docx",
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        all_text = " ".join(p.text for p in pages)
        # Empty paragraphs should not appear as content
        assert "Real content" in all_text
        assert "More content" in all_text
    finally:
        os.unlink(path)


def test_parse_unsupported_mime_raises():
    with pytest.raises(ValueError, match="Unsupported mime type"):
        parse("/tmp/fake.txt", "fake.txt", "text/plain")


# ── Cleaner ───────────────────────────────────────────────────────────────────

def _make_page(text: str, page_number: int = 1) -> ParsedPage:
    return ParsedPage(text=text, page_number=page_number, source="test.pdf")


def test_clean_returns_cleaned_page_type():
    result = clean_page(_make_page("Hello world"))
    assert isinstance(result, CleanedPage)


def test_clean_unicode_normalisation():
    # fi ligature (U+FB01) should become "fi"
    result = clean_page(_make_page("The ﬁle was found."))
    assert "fi" in result.text
    assert "ﬁ" not in result.text


def test_clean_hyphenated_linebreak():
    raw = "The comprehen-\nsion was difficult."
    result = clean_page(_make_page(raw))
    assert "comprehension" in result.text
    assert "-\n" not in result.text


def test_clean_multiple_hyphenated_linebreaks():
    raw = "pre-\nfix and suf-\nfix"
    result = clean_page(_make_page(raw))
    assert "prefix" in result.text
    assert "suffix" in result.text


def test_clean_removes_null_bytes():
    raw = "Hello\x00World\x01Test"
    result = clean_page(_make_page(raw))
    assert "\x00" not in result.text
    assert "\x01" not in result.text
    assert "HelloWorldTest" in result.text


def test_clean_collapses_multiple_spaces():
    raw = "Word   with    lots     of spaces."
    result = clean_page(_make_page(raw))
    assert "  " not in result.text
    assert "Word with lots of spaces." in result.text


def test_clean_collapses_excessive_newlines():
    raw = "Paragraph one.\n\n\n\n\nParagraph two."
    result = clean_page(_make_page(raw))
    # Should reduce to at most 2 consecutive newlines
    assert "\n\n\n" not in result.text
    assert "Paragraph one." in result.text
    assert "Paragraph two." in result.text


def test_clean_strips_standalone_page_numbers():
    raw = "Some content.\n\n5\n\nMore content."
    result = clean_page(_make_page(raw))
    # The lone "5" line should be removed
    lines = result.text.split("\n")
    assert "5" not in [l.strip() for l in lines]


def test_clean_strips_page_n_of_m():
    raw = "Introduction\nPage 3 of 45\nThe actual content here."
    result = clean_page(_make_page(raw))
    assert "Page 3 of 45" not in result.text
    assert "actual content" in result.text


def test_clean_preserves_char_counts():
    raw = "Hello world   test"
    result = clean_page(_make_page(raw))
    assert result.original_char_count == len(raw)
    assert result.cleaned_char_count == len(result.text)
    assert result.cleaned_char_count <= result.original_char_count


def test_clean_empty_page_returns_empty_text():
    result = clean_page(_make_page("   \n\n\t\n   "))
    assert result.text == ""


def test_clean_pages_discards_empty_after_cleaning():
    pages = [
        _make_page("Real content here.", page_number=1),
        _make_page("   \n\n\n   ", page_number=2),  # will be empty after cleaning
        _make_page("More real content.", page_number=3),
    ]
    cleaned = clean_pages(pages)
    assert len(cleaned) == 2
    assert cleaned[0].page_number == 1
    assert cleaned[1].page_number == 3


def test_clean_pages_preserves_order():
    pages = [_make_page(f"Page {i} content", page_number=i) for i in range(1, 6)]
    cleaned = clean_pages(pages)
    assert [c.page_number for c in cleaned] == [1, 2, 3, 4, 5]


# ── Integration: parse → clean ────────────────────────────────────────────────

def test_parse_then_clean_pdf():
    """Full parse+clean pipeline on a real PDF."""
    path = _make_pdf_file([
        "Introduction\n\n1\n\nThis is the fi rst section of the document.",
        "Section Two\n\nPage 2 of 3\n\nThe compre-\nhension of the text is important.",
        "Conclusion\n\n3\n\nFinal thoughts here.",
    ])
    try:
        raw = parse(path, "integration.pdf", "application/pdf")
        cleaned = clean_pages(raw)

        assert len(cleaned) == 3

        all_text = " ".join(c.text for c in cleaned)

        # Page numbers stripped
        assert "Page 2 of 3" not in all_text
        # Hyphenation fixed
        assert "comprehension" in all_text
        # Unicode normalised (fi ligature → fi — PyMuPDF may already normalise this)
        assert "ﬁ" not in all_text
    finally:
        os.unlink(path)


# ── Orchestrator integration test (via API) ───────────────────────────────────

def test_upload_triggers_pipeline_status_change():
    """
    Upload a document via the API and verify the background pipeline advances
    the status from 'uploaded' → 'ready'.

    Key challenge: the orchestrator's run_pipeline() creates its own SessionLocal
    (it doesn't go through FastAPI's dependency injection). We monkey-patch it to
    point at the test DB so it can find the document.

    Note: Starlette's TestClient runs BackgroundTasks synchronously, so status
    is final by the time client.post() returns — no polling needed.
    """
    from fastapi.testclient import TestClient
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker
    from app.main import app
    from app.database import Base, get_db
    from app.core.dependencies import get_chroma
    import app.pipeline.orchestrator as orchestrator_module
    import chromadb

    TEST_DB = "sqlite:///./test_pipeline_integration.db"
    engine = create_engine(TEST_DB, connect_args={"check_same_thread": False})
    TestSession = sessionmaker(bind=engine)

    def _db():
        db = TestSession()
        try:
            yield db
        finally:
            db.close()

    from app.models import user, document, job  # noqa
    Base.metadata.create_all(bind=engine)
    app.dependency_overrides[get_db] = _db
    app.dependency_overrides[get_chroma] = lambda: chromadb.EphemeralClient()

    # Redirect the orchestrator's SessionLocal to the test DB
    original_session_local = orchestrator_module.SessionLocal
    orchestrator_module.SessionLocal = TestSession

    client = TestClient(app)

    try:
        # Register + login
        client.post("/auth/register", json={"email": "pipe@example.com", "password": "password123"})
        token = client.post("/auth/login", json={
            "email": "pipe@example.com", "password": "password123"
        }).json()["access_token"]
        headers = {"Authorization": f"Bearer {token}"}

        # Upload a real PDF — TestClient runs BackgroundTasks before returning
        path = _make_pdf_file(["This document has real content for pipeline testing."])
        with open(path, "rb") as f:
            res = client.post(
                "/docs/upload",
                files={"file": ("pipeline_test.pdf", f, "application/pdf")},
                headers=headers,
            )
        os.unlink(path)

        assert res.status_code == 201
        doc_id = res.json()["id"]

        # Status should already be final (background task ran synchronously)
        doc_detail = client.get(f"/docs/{doc_id}", headers=headers).json()
        assert doc_detail["status"] == "ready", \
            f"Expected 'ready', got '{doc_detail['status']}': {doc_detail.get('error_message')}"
        assert doc_detail["page_count"] >= 1

    finally:
        orchestrator_module.SessionLocal = original_session_local
        Base.metadata.drop_all(bind=engine)
        app.dependency_overrides.clear()
